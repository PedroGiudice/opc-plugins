"""Guard de sobrescrita para os geradores .docx CMR.

Motivação (incidente angatu-priscila, 06/08/2026): o gerador regravou um
.docx que a operadora havia revisado no Word, destruindo a versão salva.
O guard torna essa colisão impossível por construção: `salvar_com_guard`
NUNCA sobrescreve um arquivo cujo conteúdo diverge da última gravação do
próprio gerador.

Mecanismo — sem estado externo, o fingerprint viaja no próprio arquivo:

1. Ao gravar, um hash do texto (parágrafos + tabelas) é embutido em
   `core_properties.comments` (campo "Comentários" das propriedades do
   documento; o Word o preserva ao salvar).
2. Ao regravar um caminho existente, o guard abre o arquivo em disco e
   recomputa o hash. Marca confere com o conteúdo -> ninguém tocou ->
   sobrescreve em silêncio. Marca ausente, arquivo ilegível ou hash
   divergente -> edição externa -> grava em nome carimbado e avisa.
3. Lock `~$` do Word presente -> documento aberto no Word -> desvia
   sempre (gravar sob Word aberto é inútil: o Word regrava por cima).

Quando o desvio acontece, o operador tem trabalho manual no arquivo
original: leia-o, incorpore a revisão e só então regenere.
"""

import hashlib
import re
from datetime import datetime
from pathlib import Path

from docx import Document

_MARCA_PREFIX = "cmr-gen:"


def _texto_documento(doc):
    partes = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                partes.append(celula.text)
    return re.sub(r"\s+", " ", "\n".join(partes)).strip()


def _fingerprint(doc):
    return hashlib.sha256(_texto_documento(doc).encode("utf-8")).hexdigest()[:16]


def _marca_e_fingerprint_do_disco(path):
    """Retorna (marca_gravada, fingerprint_atual) ou (None, None) se ilegível."""
    try:
        doc = Document(str(path))
    except Exception:
        return None, None
    comments = doc.core_properties.comments or ""
    marca = comments[len(_MARCA_PREFIX):] if comments.startswith(_MARCA_PREFIX) else None
    return marca, _fingerprint(doc)


def _word_aberto(path):
    """Detecta o owner file do Word (~$): documento aberto em alguma janela."""
    alvo = path.name
    for lock in path.parent.glob("~$*"):
        sufixo = lock.name[2:]
        if sufixo and alvo.endswith(sufixo):
            return True
    return False


def _caminho_carimbado(path):
    carimbo = datetime.now().strftime("%d.%m.%Y %Hh%M")
    candidato = path.with_name(f"{path.stem} - GERADO {carimbo}{path.suffix}")
    n = 2
    while candidato.exists():
        candidato = path.with_name(f"{path.stem} - GERADO {carimbo} ({n}){path.suffix}")
        n += 1
    return candidato


def salvar_com_guard(doc, caminho):
    """Grava `doc` em `caminho` sem nunca sobrescrever edição externa.

    Retorna o Path efetivamente gravado (o pedido, ou um desvio carimbado).
    """
    alvo = Path(caminho)
    doc.core_properties.comments = _MARCA_PREFIX + _fingerprint(doc)

    motivo = None
    if alvo.exists():
        if _word_aberto(alvo):
            motivo = "o documento esta ABERTO no Word (lock ~$ presente)"
        else:
            marca, atual = _marca_e_fingerprint_do_disco(alvo)
            if marca is None or atual is None:
                motivo = "o arquivo em disco nao foi gerado por este gerador (ou esta ilegivel)"
            elif marca != atual:
                motivo = "o conteudo em disco foi EDITADO fora do gerador desde a ultima geracao"

    if motivo:
        desvio = _caminho_carimbado(alvo)
        doc.save(str(desvio))
        print(f"[GUARD] NAO sobrescrevi '{alvo.name}': {motivo}.")
        print(f"[GUARD] A versao nova foi gravada ao lado como: '{desvio.name}'.")
        print(
            "[GUARD] O arquivo original pode conter revisao manual do operador: "
            "leia-o, incorpore as alteracoes e atualize a minuta .md antes de regenerar."
        )
        return desvio

    doc.save(str(alvo))
    print(f"Salvo: {alvo}")
    return alvo
