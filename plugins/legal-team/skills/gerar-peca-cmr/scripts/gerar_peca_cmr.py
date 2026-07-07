"""
Gerador de peças processuais no template CMR Advogados.

Usa python-docx com formatação inline (sem estilos Word).
Numeração automática via definições XML reais extraídas de contestacao-template.docx.

Uso:
    from gerar_peca_cmr import PecaCMR

    peca = PecaCMR()
    peca.enderecamento("EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA 2ª VARA CÍVEL...")
    peca.espacos(5)
    peca.processo("1234567-89.2025.8.26.0100")
    peca.espacos(3)
    peca.qualificacao("FULANO DE TAL, pessoa jurídica...", "CONTESTAÇÃO", "proposta por", "BELTRANO LTDA.")
    peca.espaco()
    peca.titulo_peca("CONTESTAÇÃO")
    peca.espaco()
    peca.corpo("nos termos do artigo 335...")
    peca.espaco()
    peca.capitulo("TEMPESTIVIDADE")
    peca.espaco()
    peca.corpo("Nos termos do artigo 335...")
    peca.corpo("De acordo com o ev. 24...")  # Sem espaco = mesmo fluxo
    peca.espaco()
    peca.subcapitulo("ILEGITIMIDADE PASSIVA DA RÉ")
    peca.espaco()
    peca.corpo("A legitimidade passiva...")
    peca.espaco()
    peca.citacao_bloco("O CDC é inaplicável...", "(STJ, AgInt no AREsp 2.132.923/SP...)")
    peca.espaco()
    peca.corpo("A similitude com o caso...")
    peca.espaco()
    peca.fecho("São Paulo", "09 de dezembro de 2025")
    peca.assinatura("Carlos Magno N. Rodrigues", "OAB/SP nº 129.021")

    peca.salvar("contestacao.docx")
"""

from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree
import copy


class PecaCMR:
    """Gera peça processual formatada no padrão CMR Advogados."""

    FONT = "Century Gothic"
    FONT_TIMBRE = "Verdana"  # fonte do papel timbrado (cabeçalho/rodapé)
    SIZE = Pt(12)
    SIZE_TITULO = Pt(14)
    SIZE_CAPITULO = Pt(13)
    SIZE_CITACAO = Pt(11)
    LINE_SPACING = 1.5
    MARGIN = Cm(2.5)
    FIRST_LINE_INDENT = Cm(4.0)
    CITACAO_INDENT = Cm(4.0)
    PEDIDO_INDENT = Cm(6.0)

    # IDs das listas de numeração — offset 10+ para não conflitar com defaults do python-docx
    NUM_CORPO = 10        # 1. 2. 3. (decimal sequencial)
    NUM_CAPITULO = 11     # I. II. III. (romano maiúsculo, bold)
    NUM_SUBCAP_PRELIM = 12  # A. B. (letra maiúscula, bold) — preliminares
    NUM_SUBCAP_MERITO = 13  # A. B. C. (letra maiúscula, bold) — mérito
    NUM_ITEM_ROMANO = 14  # i. ii. iii. (romano minúsculo)
    NUM_ITEM_LETRA = 15   # a. b. c. (letra minúscula)

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_defaults()
        self._setup_numbering()
        self._setup_header_footer()
        self._corpo_counter = 0

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = self.MARGIN
        section.bottom_margin = self.MARGIN
        section.left_margin = self.MARGIN
        section.right_margin = self.MARGIN
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(0.18)

    def _setup_defaults(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.FONT
        font.size = self.SIZE
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        rpr = style.element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            style.element.append(rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
            rfonts.set(qn(attr), self.FONT)

    def _run_timbre(self, paragraph, text, size, color=None):
        """Run do papel timbrado — fonte Verdana, com cor opcional, inline."""
        run = paragraph.add_run(text)
        run.font.name = self.FONT_TIMBRE
        run.font.size = size
        if color is not None:
            run.font.color.rgb = color
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
            rfonts.set(qn(attr), self.FONT_TIMBRE)
        return run

    def _setup_header_footer(self):
        """Cabeçalho e rodapé do papel timbrado CMR, com formatação INLINE.

        Valores extraídos do timbre real do escritório:
        - Cabeçalho: Verdana, alinhado à direita. "C. M. " em cinza claro
          (BFBFBF) + "RODRIGUES" em preto, ambos 26pt; abaixo "Advogados" 10pt.
        - Rodapé: Verdana 8pt, cinza escuro (404040), endereço alinhado à direita.

        Nada depende de estilos nomeados — cada run carrega sua própria fonte,
        tamanho e cor.
        """
        section = self.doc.sections[0]

        # --- Cabeçalho ---
        header = section.header
        header.is_linked_to_previous = False
        for p in list(header.paragraphs):
            p._element.getparent().remove(p._element)

        nome = header.add_paragraph()
        nome.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        nome.paragraph_format.space_after = Pt(0)
        nome.paragraph_format.line_spacing = 1.0
        nome.paragraph_format.right_indent = Twips(283)
        self._run_timbre(nome, "C. M. ", Pt(26), RGBColor(0xBF, 0xBF, 0xBF))
        self._run_timbre(nome, "RODRIGUES", Pt(26))

        sub = header.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sub.paragraph_format.space_after = Pt(0)
        sub.paragraph_format.line_spacing = 1.0
        sub.paragraph_format.right_indent = Twips(283)
        self._run_timbre(sub, "Advogados", Pt(10))

        # --- Rodapé ---
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)

        for linha in (
            "Alameda Santos nº. 211, 16º andar, cj. 1607",
            "São Paulo – SP 01419-000",
            "Tel.: (11) 3044 4160",
        ):
            fp = footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            fp.paragraph_format.space_after = Pt(0)
            fp.paragraph_format.line_spacing = 1.0
            self._run_timbre(fp, linha, Pt(8), RGBColor(0x40, 0x40, 0x40))

    def _setup_numbering(self):
        """Injeta definições de numeração automática no numbering.xml."""
        numbering_part = self.doc.part.numbering_part
        numbering_elem = numbering_part.numbering_definitions._numbering

        abstract_defs = [
            # abs#10 -> numId=10: decimal corpo (1. 2. 3.)
            {
                "abstractNumId": "10",
                "levels": [
                    {"ilvl": "0", "numFmt": "decimal", "lvlText": "%1.", "start": "1",
                     "ind_left": "2988", "ind_hanging": "360"},
                    {"ilvl": "1", "numFmt": "lowerLetter", "lvlText": "%2.", "start": "1",
                     "ind_left": "3708", "ind_hanging": "360"},
                    {"ilvl": "2", "numFmt": "lowerRoman", "lvlText": "%3.", "start": "1",
                     "ind_left": "4428", "ind_hanging": "180"},
                ],
            },
            # abs#11 -> numId=11: romano maiúsculo capítulos (I. II. III.)
            {
                "abstractNumId": "11",
                "levels": [
                    {"ilvl": "0", "numFmt": "upperRoman", "lvlText": "%1.", "start": "1",
                     "bold": True, "ind_left": "1440", "ind_hanging": "360"},
                    {"ilvl": "1", "numFmt": "lowerLetter", "lvlText": "%2.", "start": "1",
                     "ind_left": "2160", "ind_hanging": "360"},
                ],
            },
            # abs#12 -> numId=12: letra maiúscula subcapítulos preliminares (A. B.)
            {
                "abstractNumId": "12",
                "levels": [
                    {"ilvl": "0", "numFmt": "upperLetter", "lvlText": "%1.", "start": "1",
                     "bold": True, "ind_left": "720", "ind_hanging": "360"},
                    {"ilvl": "1", "numFmt": "lowerLetter", "lvlText": "%2.", "start": "1",
                     "ind_left": "1440", "ind_hanging": "360"},
                ],
            },
            # abs#13 -> numId=13: letra maiúscula subcapítulos mérito (A. B. C.)
            {
                "abstractNumId": "13",
                "levels": [
                    {"ilvl": "0", "numFmt": "upperLetter", "lvlText": "%1.", "start": "1",
                     "bold": True, "ind_left": "720", "ind_hanging": "360"},
                    {"ilvl": "1", "numFmt": "lowerLetter", "lvlText": "%2.", "start": "1",
                     "ind_left": "1440", "ind_hanging": "360"},
                ],
            },
            # abs#14 -> numId=14: romano minúsculo itens (i. ii. iii.)
            {
                "abstractNumId": "14",
                "levels": [
                    {"ilvl": "0", "numFmt": "lowerRoman", "lvlText": "%1.", "start": "1",
                     "ind_left": "2988", "ind_hanging": "360"},
                    {"ilvl": "1", "numFmt": "lowerLetter", "lvlText": "%2.", "start": "1",
                     "ind_left": "3708", "ind_hanging": "360"},
                ],
            },
            # abs#15 -> numId=15: letra minúscula requerimento (a. b. c.)
            {
                "abstractNumId": "15",
                "levels": [
                    {"ilvl": "0", "numFmt": "lowerLetter", "lvlText": "%1.", "start": "1",
                     "ind_left": "2628", "ind_hanging": "360"},
                    {"ilvl": "1", "numFmt": "lowerLetter", "lvlText": "%2.", "start": "1",
                     "ind_left": "3348", "ind_hanging": "360"},
                ],
            },
        ]

        num_refs = [
            {"numId": "10", "abstractNumId": "10"},
            {"numId": "11", "abstractNumId": "11"},
            {"numId": "12", "abstractNumId": "12"},
            {"numId": "13", "abstractNumId": "13"},
            {"numId": "14", "abstractNumId": "14"},
            {"numId": "15", "abstractNumId": "15"},
        ]

        for adef in abstract_defs:
            absnum = OxmlElement("w:abstractNum")
            absnum.set(qn("w:abstractNumId"), adef["abstractNumId"])

            multiline = OxmlElement("w:multiLevelType")
            multiline.set(qn("w:val"), "hybridMultilevel")
            absnum.append(multiline)

            for lvl_def in adef["levels"]:
                lvl = OxmlElement("w:lvl")
                lvl.set(qn("w:ilvl"), lvl_def["ilvl"])

                start = OxmlElement("w:start")
                start.set(qn("w:val"), lvl_def["start"])
                lvl.append(start)

                numfmt = OxmlElement("w:numFmt")
                numfmt.set(qn("w:val"), lvl_def["numFmt"])
                lvl.append(numfmt)

                lvltext = OxmlElement("w:lvlText")
                lvltext.set(qn("w:val"), lvl_def["lvlText"])
                lvl.append(lvltext)

                lvljc = OxmlElement("w:lvlJc")
                lvljc.set(qn("w:val"), "left")
                lvl.append(lvljc)

                ppr = OxmlElement("w:pPr")
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), lvl_def["ind_left"])
                ind.set(qn("w:hanging"), lvl_def["ind_hanging"])
                ppr.append(ind)
                lvl.append(ppr)

                if lvl_def.get("bold"):
                    rpr = OxmlElement("w:rPr")
                    b = OxmlElement("w:b")
                    rpr.append(b)
                    lvl.append(rpr)

                absnum.append(lvl)

            numbering_elem.append(absnum)

        for nref in num_refs:
            num = OxmlElement("w:num")
            num.set(qn("w:numId"), nref["numId"])
            absref = OxmlElement("w:abstractNumId")
            absref.set(qn("w:val"), nref["abstractNumId"])
            num.append(absref)
            numbering_elem.append(num)

    def _add_run(self, paragraph, text, bold=False, italic=False, size=None, font=None):
        run = paragraph.add_run(text)
        run.font.name = font or self.FONT
        if size:
            run.font.size = size
        run.bold = bold
        run.italic = italic
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        f = font or self.FONT
        for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
            rfonts.set(qn(attr), f)
        return run

    def _set_numbering(self, paragraph, num_id, ilvl=0):
        ppr = paragraph._element.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            paragraph._element.insert(0, ppr)
        numpr = OxmlElement("w:numPr")
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), str(ilvl))
        numpr.append(ilvl_elem)
        numid_elem = OxmlElement("w:numId")
        numid_elem.set(qn("w:val"), str(num_id))
        numpr.append(numid_elem)
        ppr.append(numpr)

    def _set_indent(self, paragraph, first_line=None, left=None, right=None):
        pf = paragraph.paragraph_format
        if first_line is not None:
            pf.first_line_indent = first_line
        if left is not None:
            pf.left_indent = left
        if right is not None:
            pf.right_indent = right

    def _p(self, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """Cria um parágrafo com o alinhamento aplicado DIRETO (jc no pPr).

        Justificação é regra no padrão de petição. Delegar isso ao estilo Normal
        (herança) falha em vários renderizadores — por isso o jc é inline em todo
        parágrafo de texto. Center só onde é título.
        """
        par = self.doc.add_paragraph()
        if align is not None:
            par.alignment = align
        return par

    def espaco(self):
        """Adiciona parágrafo vazio (Enter literal) — separador padrão."""
        self.doc.add_paragraph()

    def espacos(self, n):
        """Adiciona N parágrafos vazios."""
        for _ in range(n):
            self.doc.add_paragraph()

    def enderecamento(self, texto):
        """Endereçamento ao juízo — bold, sem recuo."""
        p = self._p()
        self._add_run(p, texto, bold=True)

    def processo(self, numero):
        """Número do processo — bold."""
        p = self._p()
        self._add_run(p, f"Processo nº {numero}", bold=True)

    def qualificacao(self, texto):
        """Qualificação das partes — recuo 4cm na primeira linha.

        Para negritar nomes de partes, use qualificacao_complexa().
        """
        p = self._p()
        self._set_indent(p, first_line=self.FIRST_LINE_INDENT)
        self._add_run(p, texto)

    def qualificacao_complexa(self, *segmentos):
        """Qualificação com formatação mista.

        Cada segmento é (texto, bold) ou (texto, bold, italic).
        Ex: peca.qualificacao_complexa(
                ("SALESFORCE TECNOLOGIA LTDA.", True),
                (", pessoa jurídica de direito privado...", False),
            )
        """
        p = self._p()
        self._set_indent(p, first_line=self.FIRST_LINE_INDENT)
        for seg in segmentos:
            text = seg[0]
            bold = seg[1] if len(seg) > 1 else False
            italic = seg[2] if len(seg) > 2 else False
            self._add_run(p, text, bold=bold, italic=italic)

    def titulo_peca(self, texto):
        """Título da peça — centralizado, 14pt, bold."""
        p = self._p(WD_ALIGN_PARAGRAPH.CENTER)
        self._add_run(p, texto, bold=True, size=self.SIZE_TITULO)

    def capitulo(self, texto):
        """Título de capítulo — I. II. III. (romano maiúsculo, 13pt, bold).

        Numeração automática.
        """
        p = self._p()
        self._set_numbering(p, self.NUM_CAPITULO, ilvl=0)
        self._set_indent(p, left=Cm(0), first_line=Cm(0.25))
        self._add_run(p, texto, bold=True, size=self.SIZE_CAPITULO)

    def subcapitulo(self, texto, num_id=None):
        """Subcapítulo — A. B. C. (letra maiúscula, bold).

        num_id: NUM_SUBCAP_PRELIM (3) para preliminares,
                NUM_SUBCAP_MERITO (4) para mérito.
                Default: NUM_SUBCAP_MERITO.
        """
        if num_id is None:
            num_id = self.NUM_SUBCAP_MERITO
        p = self._p()
        self._set_numbering(p, num_id, ilvl=0)
        self._set_indent(p, left=Cm(0), first_line=Cm(0))
        self._add_run(p, texto, bold=True)

    def corpo(self, texto):
        """Parágrafo de corpo — numeração decimal sequencial, recuo 4cm."""
        p = self._p()
        self._set_numbering(p, self.NUM_CORPO, ilvl=0)
        self._set_indent(p, left=Cm(0), first_line=self.FIRST_LINE_INDENT)
        self._add_run(p, texto)

    def corpo_complexo(self, *segmentos):
        """Parágrafo de corpo com formatação mista (bold/italic inline).

        Cada segmento é (texto,) ou (texto, bold) ou (texto, bold, italic).
        """
        p = self._p()
        self._set_numbering(p, self.NUM_CORPO, ilvl=0)
        self._set_indent(p, left=Cm(0), first_line=self.FIRST_LINE_INDENT)
        for seg in segmentos:
            if isinstance(seg, str):
                self._add_run(p, seg)
            else:
                text = seg[0]
                bold = seg[1] if len(seg) > 1 else False
                italic = seg[2] if len(seg) > 2 else False
                self._add_run(p, text, bold=bold, italic=italic)

    def corpo_sem_numero(self, texto):
        """Parágrafo sem numeração — recuo 4cm. Para continuações."""
        p = self._p()
        self._set_indent(p, first_line=self.FIRST_LINE_INDENT)
        self._add_run(p, texto)

    def citacao_bloco(self, ementa, referencia=None):
        """Citação em bloco — 11pt, itálico, recuo esquerdo 4cm.

        ementa: texto da ementa (será itálico)
        referencia: "(STJ, REsp...)" — 11pt, sem itálico
        """
        p = self._p()
        self._set_indent(p, left=self.CITACAO_INDENT)
        self._add_run(p, ementa, italic=True, size=self.SIZE_CITACAO)
        if referencia:
            self._add_run(p, f" {referencia}", size=self.SIZE_CITACAO)

    def citacao_bloco_multi(self, *paragrafos_ementa, referencia=None):
        """Citação com múltiplos parágrafos de ementa.

        Cada parágrafo_ementa é um texto que vira parágrafo separado.
        O último recebe a referência.
        """
        for i, texto in enumerate(paragrafos_ementa):
            p = self._p()
            self._set_indent(p, left=self.CITACAO_INDENT)
            self._add_run(p, texto, italic=True, size=self.SIZE_CITACAO)
            if i == len(paragrafos_ementa) - 1 and referencia:
                self._add_run(p, f" {referencia}", size=self.SIZE_CITACAO)

    def item_pedido(self, texto):
        """Item de pedido — romano minúsculo (i. ii. iii.), recuo 6cm."""
        p = self._p()
        self._set_numbering(p, self.NUM_ITEM_ROMANO, ilvl=0)
        self._set_indent(p, left=self.PEDIDO_INDENT)
        self._add_run(p, texto)

    def item_requerimento(self, texto):
        """Item de requerimento final — letra minúscula (a. b. c.)."""
        p = self._p()
        self._set_numbering(p, self.NUM_ITEM_LETRA, ilvl=0)
        self._set_indent(p, left=self.PEDIDO_INDENT)
        self._add_run(p, texto)

    def sub_item(self, texto):
        """Sub-item — letra minúscula, nível 1 da lista de corpo."""
        p = self._p()
        self._set_numbering(p, self.NUM_CORPO, ilvl=1)
        self._add_run(p, texto)

    def fecho(self, cidade, data):
        """Fecho padrão."""
        p1 = self._p()
        self._set_indent(p1, first_line=self.FIRST_LINE_INDENT)
        self._add_run(p1, "Termos em que,")
        self.espaco()
        p2 = self._p()
        self._set_indent(p2, first_line=self.FIRST_LINE_INDENT)
        self._add_run(p2, "Pede deferimento.")
        self.espaco()
        p3 = self._p()
        self._set_indent(p3, first_line=self.FIRST_LINE_INDENT)
        self._add_run(p3, f"{cidade}, {data}.")

    def assinatura(self, nome, oab):
        """Bloco de assinatura."""
        self.espacos(2)
        p1 = self._p()
        self._add_run(p1, nome)
        p2 = self._p()
        self._add_run(p2, oab)

    def assinatura_dupla(self, nome1, oab1, nome2, oab2):
        """Duas assinaturas lado a lado (simplificado: uma abaixo da outra)."""
        self.espacos(2)
        p1 = self._p()
        self._add_run(p1, nome1)
        p1b = self._p()
        self._add_run(p1b, oab1)
        self.espaco()
        p2 = self._p()
        self._add_run(p2, nome2)
        p2b = self._p()
        self._add_run(p2b, oab2)

    def _aplicar_bordas_tabela(self, table):
        """Aplica bordas (grid) DIRETO no XML da tabela.

        Bordas inline via tblBorders são autossuficientes — não dependem do
        estilo "Table Grid", coerente com o resto do gerador (formatação direta,
        nunca por estilo nomeado).
        """
        tbl_pr = table._element.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            borders.append(el)
        tbl_pr.append(borders)

    def tabela(self, headers, rows):
        """Tabela simples — Century Gothic 12pt."""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        self._aplicar_bordas_tabela(table)
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            self._add_run(p, h, bold=True)
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = table.rows[i + 1].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                self._add_run(p, val)

    def salvar(self, path):
        self.doc.save(path)
        print(f"Salvo: {path}")


def _demo():
    """Gera documento de demonstração com todos os elementos."""
    peca = PecaCMR()

    peca.enderecamento(
        "EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA 2ª VARA CÍVEL "
        "DO FORO REGIONAL DE SANTO AMARO DA COMARCA DE SÃO PAULO – SP"
    )
    peca.espacos(5)
    peca.processo("1234567-89.2025.8.26.0100")
    peca.espacos(3)
    peca.qualificacao_complexa(
        ("EMPRESA MODELO LTDA.", True),
        (", pessoa jurídica de direito privado, inscrita no CNPJ nº "
         "12.345.678/0001-90, com sede na Rua Exemplo, nº 100, São Paulo/SP, "
         "vem, por seus advogados, nos autos da ", False),
        ("AÇÃO DE INDENIZAÇÃO", True),
        (" proposta por ", False),
        ("FULANO DE TAL LTDA.", True),
        (", apresentar", False),
    )
    peca.espaco()
    peca.titulo_peca("CONTESTAÇÃO")
    peca.espaco()
    peca.corpo(
        "nos termos do artigo 335 e seguintes do Código de Processo Civil, "
        "pelos motivos de fato e de direito que adiante passa a expor."
    )
    peca.espaco()

    # Capítulo I
    peca.capitulo("TEMPESTIVIDADE")
    peca.espaco()
    peca.corpo(
        "Nos termos do artigo 335, inciso III e artigo 231, inciso V, ambos "
        "do Código de Processo Civil, o Réu poderá oferecer contestação no "
        "prazo de 15 (quinze) dias úteis."
    )
    peca.corpo(
        "De acordo com os autos, a contagem para apresentação da contestação "
        "se iniciou em 14.11.2025 e finda em 09.12.2025. A presente "
        "contestação é tempestiva."
    )
    peca.espaco()

    # Capítulo II
    peca.capitulo("PRELIMINARES")
    # Subcapítulo sem espaco após capítulo
    peca.subcapitulo("ILEGITIMIDADE PASSIVA DA RÉ", num_id=PecaCMR.NUM_SUBCAP_PRELIM)
    peca.espaco()
    peca.corpo_complexo(
        ("A legitimidade passiva ", False),
        ("ad causam", False, True),  # italic
        (" exige que a conduta do réu seja apta, ao menos em tese, "
         "a causar o dano narrado pelo autor.", False),
    )
    peca.espaco()
    peca.citacao_bloco(
        '"De acordo com a teoria da asserção, acolhida pela jurisprudência '
        'do Superior Tribunal de Justiça para a verificação das condições da '
        'ação, a legitimidade passiva depende da relação jurídica de direito '
        'material narrada na petição inicial."',
        "(STJ, REsp 1.964.337/RJ, Rel. Min. Marco Aurélio Bellizze, "
        "j. 14/06/2022)"
    )
    peca.espaco()
    peca.corpo(
        "A similitude com o caso em exame é perfeita: a Autora não identifica "
        "nenhum ato ou omissão da Ré causador do dano."
    )
    peca.espaco()

    # Capítulo III
    peca.capitulo("MÉRITO")
    peca.subcapitulo("INEXISTÊNCIA DE VÍCIO NO PRODUTO")
    peca.espaco()
    peca.corpo_complexo(
        ("O ", False),
        ("software", False, True),
        (" contratado funcionou conforme especificado. A ", False),
        ("Ré", True),
        (" cumpriu integralmente suas obrigações contratuais.", False),
    )
    peca.espaco()

    # Tabela comparativa
    peca.tabela(
        ["Elemento", "Contrato Ré", "Contrato Terceiro"],
        [
            ["Objeto", "Licenciamento SaaS", "Implementação"],
            ["Partes", "Ré e Autora", "Terceiro e Autora"],
            ["Vigência", "36 meses", "12 meses"],
        ],
    )
    peca.espaco()

    # Capítulo IV
    peca.capitulo("PEDIDOS")
    peca.espaco()
    peca.corpo(
        "Diante do exposto, a Ré requer a V. Exa.:"
    )
    peca.espaco()
    peca.item_requerimento(
        "o acolhimento da preliminar de ilegitimidade passiva, com a "
        "extinção do processo sem resolução do mérito;"
    )
    peca.item_requerimento(
        "subsidiariamente, a total improcedência dos pedidos autorais;"
    )
    peca.item_requerimento(
        "a condenação da Autora ao pagamento de custas e honorários "
        "advocatícios fixados em 20% sobre o valor da causa."
    )
    peca.espaco()

    peca.fecho("São Paulo", "09 de dezembro de 2025")
    peca.assinatura_dupla(
        "Carlos Magno N. Rodrigues", "OAB/SP nº 129.021",
        "Ana Beatriz Vianna Paoli", "OAB/SP nº 493.706",
    )

    peca.salvar("/tmp/demo_contestacao_cmr.docx")


if __name__ == "__main__":
    _demo()
