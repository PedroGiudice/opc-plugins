"""
Gerador de aditamentos trabalhistas bilingues no template Salesforce/CMR.

FIDELIDADE EXATA POR CONSTRUCAO
-------------------------------
Diferente de gerar_contrato_cmr.py (que RECRIA a formatacao com OxmlElement
inline e por isso nunca e identico a um modelo), este modulo usa o proprio
.docx de referencia como DOCUMENTO-BASE. Toda a formatacao -- estilos,
numbering.xml, header com logo, footer, margens, tema, fontTable -- e HERDADA
do template. python-docx sequer reserializa essas partes: elas sao copiadas
byte a byte do pacote original. O gerador apenas:

  - MODO FILL  : substitui placeholders preservando o run (formatacao intacta,
                 inclusive quando o placeholder cruza varios runs)
  - MODO BUILD : clona paragrafos/linhas-modelo para gerar conteudo novo no
                 mesmo padrao (deepcopy do w:p, troca so o texto)

Referencia: "Brazil - Amendment to Employment Agreement (Remote to Hybrid).docx"
Layout do template:
  - tables[0]: tabela 1x2 bilingue (corpo PT | EN lado a lado)
  - tables[1]: tabela 4x3 de assinaturas/testemunhas
  - header da 1a pagina: logo Salesforce
  - corpo: Arial 11pt (herdado do docDefault sz=22), justificado

Uso:
    from gerar_aditamento_cmr import AditamentoCMR
    a = AditamentoCMR("Brazil - Amendment ... (Remote to Hybrid) (1).docx")
    a.preencher({
        "[name of the Employee]": "Fulano de Tal",
        "[Employee name]":        "Fulano de Tal",
        "[effective date]":       "01/07/2026",
        "[inserir data]":         "01 de julho de 2026",
    })
    a.set_data("Sao Paulo, 16 de junho de 2026")
    a.escolher_dias(remoto="2 (dois) dias", presencial="3 (tres)")  # toggles OR
    a.salvar("aditamento_preenchido.docx")
"""

from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


class AditamentoCMR:
    """Aditamento bilingue gerado a partir do template real (fidelidade 1:1)."""

    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = Document(template_path)
        # tables[0] = corpo bilingue 1x2 ; tables[1] = assinaturas 4x3
        self.tbl_corpo = self.doc.tables[0] if self.doc.tables else None
        self.tbl_assin = self.doc.tables[1] if len(self.doc.tables) > 1 else None
        if self.tbl_corpo is not None:
            self.cel_pt = self.tbl_corpo.rows[0].cells[0]
            self.cel_en = self.tbl_corpo.rows[0].cells[-1]

    # ------------------------------------------------------------------ #
    # MODO FILL — substituicao preservando formatacao                    #
    # ------------------------------------------------------------------ #
    def _iter_paragraphs(self):
        """Todos os paragrafos do documento: corpo solto + dentro de tabelas."""
        for p in self.doc.paragraphs:
            yield p
        for tb in self.doc.tables:
            for row in tb.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    @staticmethod
    def _replace_in_paragraph(p, old, new):
        """Substitui `old` por `new` em um paragrafo, preservando a formatacao.

        Word fragmenta texto em multiplos runs. Estrategia: concatena os runs,
        localiza o span do placeholder e reescreve mantendo o rPr do primeiro
        run atingido; runs subsequentes do span ficam vazios (formatacao
        preservada, conteudo realocado). Cobre tanto placeholder em run unico
        quanto cruzando runs. Retorna o numero de substituicoes feitas.
        """
        runs = p.runs
        if not runs:
            return 0
        full = "".join(r.text for r in runs)
        if old not in full:
            return 0

        # mapa caractere -> indice do run
        idx_map = []
        for ri, r in enumerate(runs):
            idx_map.extend([ri] * len(r.text))

        count = 0
        # reprocessa a cada substituicao porque os offsets mudam
        while old in full:
            start = full.index(old)
            end = start + len(old)
            first_run = idx_map[start]
            last_run = idx_map[end - 1]

            # texto de cada run no span recortado
            new_texts = {ri: "" for ri in range(first_run, last_run + 1)}
            # offset inicial do primeiro run no texto completo
            run_start = idx_map.index(first_run)
            # prefixo do primeiro run (antes do placeholder) + new + sufixo do ultimo
            prefix = runs[first_run].text[: start - run_start]
            last_run_start = len(full) - sum(len(runs[ri].text) for ri in range(last_run, len(runs)))
            suffix = runs[last_run].text[end - last_run_start:]

            new_texts[first_run] = prefix + new
            if last_run != first_run:
                new_texts[last_run] = suffix
            else:
                new_texts[first_run] = prefix + new + suffix

            for ri in range(first_run, last_run + 1):
                runs[ri].text = new_texts.get(ri, "")

            count += 1
            # recomputa estado
            full = "".join(r.text for r in runs)
            idx_map = []
            for ri, r in enumerate(runs):
                idx_map.extend([ri] * len(r.text))
        return count

    def replace(self, old, new):
        """Substitui `old` por `new` em todo o documento. Retorna nº de hits."""
        total = 0
        for p in self._iter_paragraphs():
            total += self._replace_in_paragraph(p, old, new)
        return total

    def preencher(self, mapping):
        """Aplica um dict {placeholder: valor}. Retorna {placeholder: nº hits}."""
        return {k: self.replace(k, v) for k, v in mapping.items()}

    def set_data(self, texto, ano_modelo="2025"):
        """Substitui a linha de data solta 'Sao Paulo, ___ de _____ de 2025'.

        Faz match pelo padrao do template (com sublinhados); `texto` e a linha
        final desejada, ex: 'Sao Paulo, 16 de junho de 2026'.
        """
        alvo = f"São Paulo, ___ de _____ de {ano_modelo}"
        n = self.replace(alvo, texto)
        if n == 0:
            # fallback: troca so os sublinhados se a linha tiver outro formato
            self.replace("___ de _____", texto)
        return n

    def escolher(self, placeholder_or, valor):
        """Resolve um toggle 'A OR B' substituindo o placeholder inteiro.

        Ex: escolher('[1 (um) dia OR 2 (dos) dias]', '2 (dois) dias').
        """
        return self.replace(placeholder_or, valor)

    # ------------------------------------------------------------------ #
    # MODO BUILD — clonagem de elementos-modelo (peca nova no padrao)     #
    # ------------------------------------------------------------------ #
    @staticmethod
    def _set_par_text(par, texto):
        """Troca o texto de um paragrafo mantendo o rPr do 1o run; zera o resto."""
        runs = par.runs
        if not runs:
            par.add_run(texto)
            return par
        runs[0].text = texto
        for r in runs[1:]:
            r.text = ""
        return par

    def clonar_paragrafo(self, modelo_par, novo_texto, depois_de=None):
        """Clona um w:p-modelo (preservando pPr/rPr exatos), troca o texto e
        insere apos `depois_de` (Paragraph) ou ao fim do mesmo pai do modelo.

        Use modelos extraidos do proprio template para garantir fidelidade:
        titulo, considerando, clausula numerada, fecho, etc.
        """
        novo_el = deepcopy(modelo_par._p)
        if depois_de is not None:
            depois_de._p.addnext(novo_el)
        else:
            modelo_par._p.addnext(novo_el)
        novo_par = Paragraph(novo_el, modelo_par._parent)
        self._set_par_text(novo_par, novo_texto)
        return novo_par

    def modelo(self, indice_pt):
        """Retorna o Paragraph de indice `indice_pt` na celula PT (para clonar)."""
        return self.cel_pt.paragraphs[indice_pt]

    # ------------------------------------------------------------------ #
    def salvar(self, path):
        self.doc.save(path)
        print(f"Salvo: {path}")
        return path


if __name__ == "__main__":
    import sys
    tpl = sys.argv[1] if len(sys.argv) > 1 else \
        "ref-amendment-remote-to-hybrid.docx"
    a = AditamentoCMR(tpl)
    print(f"Tabelas: {len(a.doc.tables)} | "
          f"corpo {a.tbl_corpo.rows[0].cells[0].paragraphs.__len__()} pars PT")
    a.preencher({
        "[name of the Employee]": "FULANO DE TAL",
        "[Employee name]": "FULANO DE TAL",
        "[effective date]": "01/07/2026",
        "[inserir data]": "01 de julho de 2026",
        "[inserir]": "30 dias",
        "[insert]": "30 days",
    })
    a.escolher("[2 (dos) dias OR 1 (um) dia]", "2 (dois) dias")
    a.escolher("[two (2) days OR one (1) day]", "two (2) days")
    a.set_data("São Paulo, 16 de junho de 2026")
    a.salvar("/tmp/cmr-contrato/_out_preenchido.docx")
