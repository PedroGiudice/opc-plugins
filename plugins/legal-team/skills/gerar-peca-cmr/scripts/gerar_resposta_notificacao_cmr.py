#!/usr/bin/env python3
"""Gera resposta a notificação extrajudicial no padrão CMR Advogados.

Formatação extraída de 2 modelos reais protocolados (jul/2026, cartas
Salesforce: PH Brasil 20.07 e Clube Candeias 22.07). NÃO segue o padrão das
peças processuais (PecaCMR): carta usa Arial 12pt (não Century Gothic),
espaçamento 1,15 (não 1,5), página Letter, margens 3/2/4/2cm, corpo SEM
numeração de parágrafos e quase sem ênfase tipográfica. O papel timbrado
(Verdana, C. M. RODRIGUES) é o mesmo das peças, com o ajuste do modelo real
de carta: header sem recuo à direita, footer com recuo 283 twips.

Divergências entre os 2 modelos, parametrizadas:
  - alinhamento do bloco de assinatura: centro (Candeias, mais recente,
    default) ou esquerda (PH Brasil) — assinatura(..., centralizada=False)
  - recuo da linha Ref.: por decisão do CEO (23/07/2026) o default é SEM
    recuo, alinhada à margem como o corpo — os modelos recuavam (5,5/1,5 e
    7,0/2,0cm), recuperável via ref(..., left=Cm(5.5), hanging=Cm(1.5))

Uso:
    from gerar_resposta_notificacao_cmr import RespostaNotificacaoCMR

    carta = RespostaNotificacaoCMR()
    carta.data("São Paulo", "20 de julho de 2026")
    carta.espaco()
    carta.destinatario(
        "À",
        "PH BRASIL PROMOTORA DE VENDAS LTDA.",
        "A/C Sr. Paulo Henrique Camargo Costa",
        "Via e-mail: paulo.costa@phng.com.br",
    )
    carta.espaco()
    carta.ref("Resposta à Notificação Extrajudicial de 08.07.2026 — ...")
    carta.espaco()
    carta.vocativo()
    carta.espaco()
    carta.corpo("Na qualidade de advogados da ...")
    carta.espaco()
    carta.corpo("...")
    carta.espaco()
    carta.corpo("Sem mais para o momento,")
    carta.espaco()
    carta.assinatura("Carlos Magno N. Rodrigues", "OAB/SP 129.021")
    carta.salvar("resposta.docx")

Como nos demais geradores CMR: formatação DIRETA (inline), nunca por estilo
nomeado; parágrafo vazio explícito (espaco()) entre blocos de conteúdo.
"""

from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class RespostaNotificacaoCMR:
    """Gera carta de resposta a notificação extrajudicial no padrão CMR."""

    FONT = "Arial"
    FONT_TIMBRE = "Verdana"  # papel timbrado (cabeçalho/rodapé)
    SIZE = Pt(12)
    LINE_SPACING = 1.15  # w:line=276 auto — carta NÃO usa o 1,5 das peças

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_defaults()
        self._setup_header_footer()

    # ------------------------------------------------------------------ setup

    def _setup_page(self):
        """Página e margens — valores em twips idênticos nos 2 modelos reais."""
        section = self.doc.sections[0]
        section.page_width = Twips(12246)   # Letter (21,6cm)
        section.page_height = Twips(15817)  # 27,9cm
        section.top_margin = Twips(2268)    # 4,0cm
        section.bottom_margin = Twips(1134)  # 2,0cm
        section.left_margin = Twips(1701)   # 3,0cm
        section.right_margin = Twips(1134)  # 2,0cm
        section.header_distance = Twips(720)
        section.footer_distance = Twips(720)

    def _setup_defaults(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.FONT
        font.size = self.SIZE
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        rpr = style.element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            style.element.append(rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
            rfonts.set(qn(attr), self.FONT)

    def _run_timbre(self, paragraph, text, size, color=None):
        """Run do papel timbrado — Verdana, cor opcional, inline."""
        run = paragraph.add_run(text)
        run.font.name = self.FONT_TIMBRE
        run.font.size = size
        if color is not None:
            run.font.color.rgb = color
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
            rfonts.set(qn(attr), self.FONT_TIMBRE)
        return run

    def _setup_header_footer(self):
        """Papel timbrado CMR, formatação inline (igual às cartas reais).

        - Cabeçalho: Verdana à direita, sem recuo. "C. M. " cinza claro
          (BFBFBF) + "RODRIGUES" preto, 26pt; abaixo "Advogados" 10pt.
        - Rodapé: Verdana 8pt cinza escuro (404040), à direita, recuo
          direito 283 twips.
        """
        section = self.doc.sections[0]

        header = section.header
        header.is_linked_to_previous = False
        for p in list(header.paragraphs):
            p._element.getparent().remove(p._element)

        nome = header.add_paragraph()
        nome.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        nome.paragraph_format.space_after = Pt(0)
        nome.paragraph_format.line_spacing = 1.0
        self._run_timbre(nome, "C. M. ", Pt(26), RGBColor(0xBF, 0xBF, 0xBF))
        self._run_timbre(nome, "RODRIGUES", Pt(26))

        sub = header.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sub.paragraph_format.space_after = Pt(0)
        sub.paragraph_format.line_spacing = 1.0
        self._run_timbre(sub, "Advogados", Pt(10))

        footer = section.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)

        for linha in (
            "Alameda Santos nº. 211, 16º andar, cj. 1607",
            "São Paulo – SP 01419-000",
            "Tel.: (11) 3044 4160",
        ):
            fp = footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            fp.paragraph_format.space_after = Pt(0)
            fp.paragraph_format.line_spacing = 1.0
            fp.paragraph_format.right_indent = Twips(283)
            self._run_timbre(fp, linha, Pt(8), RGBColor(0x40, 0x40, 0x40))

    # -------------------------------------------------------------- primitivas

    def _add_run(self, paragraph, text, bold=False, italic=False, size=None):
        run = paragraph.add_run(text)
        run.font.name = self.FONT
        if size:
            run.font.size = size
        run.bold = bold
        run.italic = italic
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
            rfonts.set(qn(attr), self.FONT)
        return run

    def _p(self, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """Parágrafo com alinhamento aplicado DIRETO (jc inline no pPr)."""
        par = self.doc.add_paragraph()
        if align is not None:
            par.alignment = align
        return par

    def espaco(self):
        """Parágrafo vazio (Enter literal) — separador padrão."""
        self.doc.add_paragraph()

    def espacos(self, n):
        for _ in range(n):
            self.doc.add_paragraph()

    # ------------------------------------------------------------------ blocos

    def data(self, cidade, data_extenso):
        """Linha de data — alinhada à direita. Ex: 'São Paulo, 20 de julho de 2026.'"""
        p = self._p(WD_ALIGN_PARAGRAPH.RIGHT)
        self._add_run(p, f"{cidade}, {data_extenso}.")

    def destinatario(self, vocativo, nome, *linhas):
        """Bloco do destinatário — linhas contíguas, à esquerda.

        vocativo: 'À' ou 'Ao' (concordando com o destinatário)
        nome: razão social/nome — BOLD
        linhas: endereço (opcional), 'A/C ...', 'Via e-mail: ...'
        """
        p = self._p(WD_ALIGN_PARAGRAPH.LEFT)
        self._add_run(p, vocativo)
        p = self._p(WD_ALIGN_PARAGRAPH.LEFT)
        self._add_run(p, nome, bold=True)
        for linha in linhas:
            p = self._p(WD_ALIGN_PARAGRAPH.LEFT)
            self._add_run(p, linha)

    def ref(self, texto, prefixo="Ref.: ", left=None, hanging=None):
        """Linha de referência — alinhada à margem, como o corpo (sem recuo).

        texto: assunto SEM o prefixo (o 'Ref.: ' é adicionado aqui).
        left/hanging (Cm) reproduzem o recuo dos modelos antigos, se pedido.
        """
        p = self._p(WD_ALIGN_PARAGRAPH.LEFT)
        pf = p.paragraph_format
        if left is not None:
            pf.left_indent = left
        if hanging is not None:
            pf.first_line_indent = -hanging
        self._add_run(p, f"{prefixo}{texto}")

    def vocativo(self, texto="Prezados Senhores,"):
        p = self._p(WD_ALIGN_PARAGRAPH.LEFT)
        self._add_run(p, texto)

    def corpo(self, texto):
        """Parágrafo de corpo — justificado, SEM numeração, SEM recuo."""
        p = self._p()
        self._add_run(p, texto)

    def corpo_complexo(self, *segmentos):
        """Parágrafo de corpo com formatação mista (raro em carta).

        Cada segmento é str, (texto, bold) ou (texto, bold, italic).
        """
        p = self._p()
        for seg in segmentos:
            if isinstance(seg, str):
                self._add_run(p, seg)
            else:
                text = seg[0]
                bold = seg[1] if len(seg) > 1 else False
                italic = seg[2] if len(seg) > 2 else False
                self._add_run(p, text, bold=bold, italic=italic)

    def assinatura(self, nome, oab, centralizada=True, saudacao="Atenciosamente,"):
        """Bloco final: saudação + 3 vazios (espaço de assinatura) + nome + OAB.

        centralizada=True (Candeias, mais recente) ou False = à esquerda (PH).
        Chamar após o último parágrafo de corpo + espaco().
        """
        align = WD_ALIGN_PARAGRAPH.CENTER if centralizada else WD_ALIGN_PARAGRAPH.LEFT
        p = self._p(align)
        self._add_run(p, saudacao)
        self.espacos(3)
        p = self._p(align)
        self._add_run(p, nome)
        p = self._p(align)
        self._add_run(p, oab)

    def salvar(self, path):
        self.doc.save(path)
        return path
