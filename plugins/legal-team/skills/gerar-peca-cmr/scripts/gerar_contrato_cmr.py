"""
Gerador de contratos societários no template CMR Advogados.

Usa python-docx com formatação inline. Padrão extraído de:
- Minuta Contrato Social OAS CMR 29 04 2026 V2.docx
- OAS Minuta Acordo de Quotistas v CMR 28 04 2026 V2.docx

Uso:
    from gerar_contrato_cmr import ContratoCMR
    c = ContratoCMR()
    c.titulo("CONTRATO SOCIAL DE CONSTITUIÇÃO DA EMPRESA LTDA.")
    c.preambulo("Pelo presente instrumento particular, os abaixo assinados...")
    c.qualificacao("[QUALIFICAÇÃO COMPLETA DA SÓCIA 1 — PJ];")
    c.clausula("1ª", "DENOMINAÇÃO, SEDE, OBJETO E DURAÇÃO")
    c.artigo("1º", "Definições")  # Para acordos de quotistas
    c.subartigo("1.1.", "Texto do subartigo...")
    c.alinea("texto da alínea")  # (a) (b) (c) automático
    c.paragrafo("único", "Texto do parágrafo...")
    c.salvar("contrato.docx")
"""

from docx import Document
from docx.shared import Pt, Cm, Mm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ContratoCMR:
    """Gera contrato societário formatado no padrão CMR Advogados."""

    FONT = "Arial"
    SIZE = Pt(12)
    LINE_SPACING_VAL = 276  # 1.15x (276/240)
    MARGIN_TOP = Mm(40)
    MARGIN_BOTTOM = Mm(20)
    MARGIN_LEFT = Mm(30)
    MARGIN_RIGHT = Mm(20)

    # IDs de numeração — offset 20+ para não conflitar
    NUM_ALINEA = 20  # (a) (b) (c) — lowerLetter com parênteses

    def __init__(self, margin_top=None):
        self.doc = Document()
        if margin_top is not None:
            self.MARGIN_TOP = margin_top
        self._setup_page()
        self._setup_defaults()
        self._setup_numbering()
        self._setup_footer()

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Mm(216)
        section.page_height = Mm(279)
        section.top_margin = self.MARGIN_TOP
        section.bottom_margin = self.MARGIN_BOTTOM
        section.left_margin = self.MARGIN_LEFT
        section.right_margin = self.MARGIN_RIGHT

    def _setup_defaults(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.FONT
        font.size = self.SIZE
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        # Set line spacing via XML (276 twips = 1.15x)
        ppr = style.element.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            style.element.append(ppr)
        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            ppr.append(spacing)
        spacing.set(qn("w:line"), str(self.LINE_SPACING_VAL))
        spacing.set(qn("w:lineRule"), "auto")

        # Font fallback
        rpr = style.element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            style.element.append(rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
            rfonts.set(qn(attr), self.FONT)

    def _setup_footer(self):
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page number field
        run = fp.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._element.append(fld_char_begin)

        run2 = fp.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run2._element.append(instr)

        run3 = fp.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run3._element.append(fld_char_end)

    def _setup_numbering(self):
        numbering_part = self.doc.part.numbering_part
        numbering_elem = numbering_part.numbering_definitions._numbering

        # abs#20: (a) (b) (c) — lowerLetter com parênteses
        absnum = OxmlElement("w:abstractNum")
        absnum.set(qn("w:abstractNumId"), "20")

        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "hybridMultilevel")
        absnum.append(multi)

        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        numfmt = OxmlElement("w:numFmt")
        numfmt.set(qn("w:val"), "lowerLetter")
        lvl.append(numfmt)

        lvltext = OxmlElement("w:lvlText")
        lvltext.set(qn("w:val"), "(%1)")
        lvl.append(lvltext)

        lvljc = OxmlElement("w:lvlJc")
        lvljc.set(qn("w:val"), "left")
        lvl.append(lvljc)

        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        ppr.append(ind)
        lvl.append(ppr)

        absnum.append(lvl)
        numbering_elem.append(absnum)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "20")
        absref = OxmlElement("w:abstractNumId")
        absref.set(qn("w:val"), "20")
        num.append(absref)
        numbering_elem.append(num)

    def _add_run(self, paragraph, text, bold=False, italic=False, size=None):
        run = paragraph.add_run(text)
        run.font.name = self.FONT
        if size:
            run.font.size = size
        run.bold = bold
        run.italic = italic
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
            rfonts.set(qn(attr), self.FONT)
        return run

    def _set_numbering(self, paragraph, num_id, ilvl=0):
        ppr = paragraph._element.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            paragraph._element.insert(0, ppr)
        numpr = OxmlElement("w:numPr")
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), str(ilvl))
        numpr.append(ilvl_elem)
        numid_elem = OxmlElement("w:numId")
        numid_elem.set(qn("w:val"), str(num_id))
        numpr.append(numid_elem)
        ppr.append(numpr)

    def espaco(self):
        self.doc.add_paragraph()

    def espacos(self, n):
        for _ in range(n):
            self.doc.add_paragraph()

    def titulo(self, texto):
        """Título do contrato — centralizado, bold."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_run(p, texto, bold=True)

    def preambulo(self, texto):
        """Texto introdutório do contrato."""
        p = self.doc.add_paragraph()
        self._add_run(p, texto)

    def qualificacao(self, texto):
        """Qualificação de sócio/parte."""
        p = self.doc.add_paragraph()
        self._add_run(p, texto)

    def qualificacao_complexa(self, *segmentos):
        """Qualificação com formatação mista — (texto, bold) ou (texto, bold, italic)."""
        p = self.doc.add_paragraph()
        for seg in segmentos:
            text = seg[0]
            bold = seg[1] if len(seg) > 1 else False
            italic = seg[2] if len(seg) > 2 else False
            self._add_run(p, text, bold=bold, italic=italic)

    def clausula(self, numero, titulo_texto):
        """Cláusula — "CLÁUSULA Xª — TÍTULO" (bold)."""
        p = self.doc.add_paragraph()
        self._add_run(p, f"CLÁUSULA {numero} — {titulo_texto}", bold=True)

    def artigo(self, numero, titulo_texto):
        """Artigo — "Artigo Xº — Título" (bold). Para acordos de quotistas."""
        p = self.doc.add_paragraph()
        self._add_run(p, f"Artigo {numero} — {titulo_texto}", bold=True)

    def capitulo(self, numero_romano, titulo_texto):
        """Capítulo — "CAPÍTULO I — TÍTULO" (centralizado, bold). Para acordos."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_run(p, f"CAPÍTULO {numero_romano} — {titulo_texto}", bold=True)

    def subartigo(self, numero, texto):
        """Subartigo — "X.Y. Texto" (bold)."""
        p = self.doc.add_paragraph()
        self._add_run(p, f"{numero} {texto}", bold=True)

    def corpo(self, texto):
        """Parágrafo de corpo — sem numeração, sem recuo especial."""
        p = self.doc.add_paragraph()
        self._add_run(p, texto)

    def corpo_complexo(self, *segmentos):
        """Parágrafo com formatação mista."""
        p = self.doc.add_paragraph()
        for seg in segmentos:
            if isinstance(seg, str):
                self._add_run(p, seg)
            else:
                text = seg[0]
                bold = seg[1] if len(seg) > 1 else False
                italic = seg[2] if len(seg) > 2 else False
                self._add_run(p, text, bold=bold, italic=italic)

    def paragrafo(self, numero, texto):
        """Parágrafo numerado — "Parágrafo Xº — texto" ou "Parágrafo único — texto"."""
        p = self.doc.add_paragraph()
        if numero.lower() == "único" or numero.lower() == "unico":
            self._add_run(p, f"Parágrafo único — {texto}")
        else:
            self._add_run(p, f"Parágrafo {numero} — {texto}")

    def alinea(self, texto):
        """Alínea — (a) (b) (c) numeração automática."""
        p = self.doc.add_paragraph()
        self._set_numbering(p, self.NUM_ALINEA, ilvl=0)
        self._add_run(p, texto)

    def item_romano(self, numero, texto):
        """Item com romano — "I — Texto" (bold). Para matérias reservadas etc."""
        p = self.doc.add_paragraph()
        self._add_run(p, f"{numero} — {texto}", bold=True)

    def definicao(self, termo, significado):
        """Definição — "Termo" em bold + significado. Para seção de definições."""
        p = self.doc.add_paragraph()
        self._set_numbering(p, self.NUM_ALINEA, ilvl=0)
        self._add_run(p, f"\"{termo}\"", bold=True)
        self._add_run(p, f" significa {significado}")

    def fecho(self, cidade, data):
        """Fecho do contrato."""
        p = self.doc.add_paragraph()
        self._add_run(p, f"{cidade}, {data}.")

    def assinatura(self, nome, cargo=None):
        """Bloco de assinatura."""
        self.espacos(2)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_run(p, "_" * 40)
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_run(p2, nome, bold=True)
        if cargo:
            p3 = self.doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_run(p3, cargo)

    def testemunhas(self, nome1, cpf1, nome2, cpf2):
        """Bloco de testemunhas."""
        self.espacos(2)
        p = self.doc.add_paragraph()
        self._add_run(p, "Testemunhas:", bold=True)
        self.espaco()
        p1 = self.doc.add_paragraph()
        self._add_run(p1, "_" * 40)
        p1b = self.doc.add_paragraph()
        self._add_run(p1b, f"Nome: {nome1}")
        p1c = self.doc.add_paragraph()
        self._add_run(p1c, f"CPF: {cpf1}")
        self.espaco()
        p2 = self.doc.add_paragraph()
        self._add_run(p2, "_" * 40)
        p2b = self.doc.add_paragraph()
        self._add_run(p2b, f"Nome: {nome2}")
        p2c = self.doc.add_paragraph()
        self._add_run(p2c, f"CPF: {cpf2}")

    def salvar(self, path):
        self.doc.save(path)
        print(f"Salvo: {path}")


def _demo():
    """Gera contrato social de demonstração."""
    c = ContratoCMR()

    c.titulo("CONTRATO SOCIAL DE CONSTITUIÇÃO DA EMPRESA MODELO LTDA.")
    c.espaco()
    c.preambulo(
        "Pelo presente instrumento particular, os abaixo assinados, a saber:"
    )
    c.espaco()
    c.qualificacao("[QUALIFICAÇÃO COMPLETA DA SÓCIA 1 — PJ];")
    c.espaco()
    c.qualificacao("[QUALIFICAÇÃO COMPLETA DA SÓCIA 2 — PJ]; e")
    c.espaco()
    c.qualificacao("[QUALIFICAÇÃO COMPLETA DA SÓCIA 3 — PJ],")
    c.espaco()
    c.corpo(
        "têm entre si, justo e contratado, constituir uma sociedade limitada, "
        "que se regerá pelas cláusulas e condições seguintes:"
    )

    c.clausula("1ª", "DENOMINAÇÃO, SEDE, OBJETO E DURAÇÃO")
    c.espaco()
    c.corpo("A Sociedade tem a denominação de EMPRESA MODELO LTDA.")
    c.espaco()
    c.corpo(
        "A Sociedade tem sede na cidade de São Paulo, Estado de São Paulo, "
        "na Alameda Santos, 211, 16º andar."
    )
    c.espaco()
    c.paragrafo("único",
        "A Sociedade poderá abrir e fechar filiais, subsidiárias, agências "
        "e escritórios em qualquer localidade do país ou do exterior."
    )
    c.espaco()
    c.corpo("A Sociedade tem por objeto a importação e comércio de cosméticos.")
    c.espaco()
    c.corpo("A Sociedade tem prazo de duração indeterminado.")
    c.espaco()

    c.clausula("2ª", "CAPITAL SOCIAL E QUOTAS")
    c.espaco()
    c.corpo(
        "O capital da Sociedade, totalmente integralizado em moeda corrente "
        "nacional, é de R$ 100.000,00, dividido em 100.000 quotas, no valor "
        "nominal de R$ 1,00 cada, assim distribuídas:"
    )
    c.espaco()
    c.alinea("[SÓCIA 1] possui 40.000 quotas, no valor total de R$ 40.000,00;")
    c.espaco()
    c.alinea("[SÓCIA 2] possui 30.000 quotas, no valor total de R$ 30.000,00; e")
    c.espaco()
    c.alinea("[SÓCIA 3] possui 30.000 quotas, no valor total de R$ 30.000,00.")
    c.espaco()
    c.paragrafo("1º",
        "A responsabilidade de cada sócia é restrita ao valor de suas quotas."
    )
    c.espaco()
    c.paragrafo("2º",
        "As sócias não respondem pelas obrigações sociais, nem mesmo "
        "subsidiariamente, ressalvadas as hipóteses legais."
    )
    c.espaco()

    c.clausula("3ª", "ADMINISTRAÇÃO")
    c.espaco()
    c.corpo(
        "A administração da Sociedade será exercida por administradora "
        "nomeada pelas sócias em reunião."
    )
    c.espaco()
    c.paragrafo("1º",
        "Os seguintes atos somente poderão ser praticados com autorização "
        "prévia de sócias representando, no mínimo, 3/4 do capital social:"
    )
    c.espaco()
    c.alinea(
        "a compra, venda, locação ou hipoteca de imóveis da Sociedade;"
    )
    c.espaco()
    c.alinea("qualquer alteração no objeto social da Sociedade;")
    c.espaco()
    c.alinea(
        "a participação em outras sociedades, como acionista ou sócia; e"
    )
    c.espaco()
    c.alinea(
        "qualquer outro ato não compreendido no curso regular dos negócios."
    )
    c.espaco()

    c.clausula("4ª", "EXERCÍCIO SOCIAL")
    c.espaco()
    c.corpo(
        "O exercício social terminará no dia 31 de dezembro de cada ano."
    )
    c.espaco()

    c.clausula("5ª", "FORO")
    c.espaco()
    c.corpo(
        "Fica eleito o Foro da Comarca de São Paulo, Estado de São Paulo, "
        "para dirimir quaisquer questões oriundas deste contrato."
    )
    c.espaco()

    c.fecho("São Paulo", "[DATA]")
    c.assinatura("[SÓCIA 1]", "Representante Legal")
    c.assinatura("[SÓCIA 2]", "Representante Legal")
    c.assinatura("[SÓCIA 3]", "Representante Legal")
    c.testemunhas("[NOME]", "[CPF]", "[NOME]", "[CPF]")

    c.salvar("/tmp/demo_contrato_cmr.docx")


if __name__ == "__main__":
    _demo()
